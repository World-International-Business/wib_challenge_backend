from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from challenges.models import Submission, Challenge
from questions.models import Domain, Question


def generate_evaluations_report(submissions, output_path):
    """
    Génère un rapport Word pour une liste d'évaluations.
    
    Args:
        submissions: Liste des objets Submission
        output_path: Chemin de sortie pour le document Word
    """
    # Créer un document Word
    doc = Document()
    
    # Configuration de base du document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Ajouter un titre principal
    title = doc.add_heading('Rapport d\'évaluations', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Pour chaque évaluation
    for submission in submissions:
        # Ajouter les informations du candidat
        doc.add_heading(f'Candidat: {submission.candidate.get_full_name()}', level=1)
        
        # Table d'informations sur le candidat
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = True
        
        # Première ligne avec email et date
        row_cells = table.rows[0].cells
        row_cells[0].text = f'Email: {submission.candidate.email}'
        row_cells[1].text = f'Date: {submission.submitted_at.strftime("%d/%m/%Y %H:%M")}'
        
        # Deuxième ligne avec challenge et domaine
        row = table.add_row().cells
        row[0].text = f'Challenge: {submission.challenge.title}'
        row[1].text = f'Domaine: {submission.challenge.domain.name}'
        
        # Troisième ligne avec résultat
        row = table.add_row().cells
        row[0].text = f'Résultat: {submission.result_percent:.2f}%'
        row[1].text = f'Status: {submission.get_status_display()}'
        
        doc.add_paragraph()  # Espacement
        
        # Ajouter les questions et réponses
        doc.add_heading('Questions et réponses', level=2)
        
        for i, answer in enumerate(submission.answers.all(), 1):
            # Titre de la question
            q_title = doc.add_paragraph()
            q_title.add_run(f'Question {i}: {answer.question.title}').bold = True
            
            # Informations sur la question
            q_info = doc.add_paragraph()
            q_info.add_run(f'Catégorie: {answer.question.category.name}\n')
            q_info.add_run(f'Type: {answer.question.get_question_type_display()}\n')
            q_info.add_run(f'Niveau: {answer.question.get_level_display()}\n')
            
            # Section réponse
            response_p = doc.add_paragraph()
            response_p.add_run('Réponse: ').bold = True
            
            if answer.question.is_open_answer:
                response_p.add_run(f'{answer.text}\n')
                
                # Afficher si la réponse est correcte
                result_p = doc.add_paragraph()
                result_run = result_p.add_run(f'Évaluation: {"Correct" if answer.is_correct else "Incorrect"}')
                result_run.bold = True
                result_run.font.color.rgb = RGBColor(0, 128, 0) if answer.is_correct else RGBColor(255, 0, 0)
                
            else:  # Questions à choix
                # Tableau des choix
                choices_table = doc.add_table(rows=1, cols=3)
                choices_table.style = 'Table Grid'
                choices_header = choices_table.rows[0].cells
                choices_header[0].text = "Choix"
                choices_header[1].text = "Sélectionné"
                choices_header[2].text = "Correct"
                
                # Remplir les choix
                for choice in answer.question.choices.all():
                    selected = choice in answer.selected_choices.all()
                    row = choices_table.add_row().cells
                    row[0].text = choice.text
                    row[1].text = "✓" if selected else ""
                    row[2].text = "✓" if choice.is_correct else ""
            
            doc.add_paragraph()  # Espacement entre questions
        
        # Séparateur entre candidats
        doc.add_page_break()
    
    # Enregistrer le document
    doc.save(output_path)
    return output_path


def generate_rh_consultant_report():
    """
    Génère un rapport pour les évaluations du domaine Consultant RH
    """
    try:
        # Récupérer le domaine Consultant RH
        rh_domain = Domain.objects.get(name="Consultant RH")
        
        # Récupérer les challenges associés à ce domaine
        rh_challenges = Challenge.objects.filter(domain=rh_domain)
        
        # Récupérer les soumissions pour ces challenges
        submissions = Submission.objects.filter(
            challenge__in=rh_challenges,
            status=Submission.CorrectionStatus.CORRECTED
        ).order_by('-submitted_at')
        
        if not submissions:
            print("Aucune évaluation trouvée pour le domaine Consultant RH")
            return None
            
        # Générer le rapport
        output_path = 'rapport_evaluations_consultant_rh.docx'
        return generate_evaluations_report(submissions, output_path)
        
    except Domain.DoesNotExist:
        print("Le domaine 'Consultant RH' n'existe pas")
        return None
    except Exception as e:
        print(f"Erreur lors de la génération du rapport: {str(e)}")
        return None


if __name__ == "__main__":
    # Ce bloc s'exécute uniquement si le fichier est lancé directement
    import django
    import os
    os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'wib_challenge.settings')
    django.setup()
    
    # Tester la génération de rapport pour Consultant RH
    output = generate_rh_consultant_report()
    if output:
        print(f"Rapport généré avec succès: {output}")
