import os
from django.core.management.base import BaseCommand
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from challenges.models import Challenge, Submission
from questions.models import Domain


class Command(BaseCommand):
    help = 'Génère des rapports Word pour les challenges'

    def add_arguments(self, parser):
        parser.add_argument('--challenge-id', type=int, nargs='+', help='IDs des challenges')
        parser.add_argument('--domain', type=str, help='Nom du domaine pour filtrer les challenges')
        parser.add_argument('--output-dir', type=str, default='reports', help='Répertoire de sortie')

    def handle(self, *args, **options):
        output_dir = options['output_dir']
        os.makedirs(output_dir, exist_ok=True)

        if options['domain']:
            try:
                domain = Domain.objects.get(name=options['domain'])
                challenges = Challenge.objects.filter(domain=domain)
                self.stdout.write(f"Génération de rapports pour le domaine {domain.name}")
            except Domain.DoesNotExist:
                self.stderr.write(f"Domaine '{options['domain']}' non trouvé")
                return
        elif options['challenge_id']:
            challenges = Challenge.objects.filter(id__in=options['challenge_id'])
            self.stdout.write(f"Génération de rapports pour {len(challenges)} challenge(s)")
        else:
            challenges = Challenge.objects.all()
            self.stdout.write(f"Génération de rapports pour tous les challenges ({challenges.count()})")

        for challenge in challenges:
            submissions = Submission.objects.filter(
                challenge=challenge,
                status=Submission.CorrectionStatus.CORRECTED
            ).order_by('-submitted_at')

            if not submissions:
                self.stdout.write(f"Aucune soumission trouvée pour le challenge '{challenge.title}'")
                continue

            output_path = os.path.join(output_dir, f"rapport_{challenge.slug}.docx")
            self.generate_challenge_report(challenge, submissions, output_path)
            self.stdout.write(f"Rapport généré : {output_path}")

        # Test spécifique pour le domaine Consultant RH
        if options['domain'] == 'Consultant RH' or not options['domain']:
            self.generate_consultant_rh_report(output_dir)

    def generate_consultant_rh_report(self, output_dir):
        """Génère un rapport pour le domaine Consultant RH"""
        try:
            domain = Domain.objects.get(name="Consultant RH")
            challenges = Challenge.objects.filter(domain=domain)

            if not challenges:
                self.stdout.write("Aucun challenge trouvé pour le domaine Consultant RH")
                return

            all_submissions = []
            for challenge in challenges:
                submissions = Submission.objects.filter(
                    challenge=challenge,
                    status=Submission.CorrectionStatus.CORRECTED
                ).order_by('-submitted_at')
                all_submissions.extend(submissions)

            if not all_submissions:
                self.stdout.write("Aucune soumission trouvée pour le domaine Consultant RH")
                return

            output_path = os.path.join(output_dir, "rapport_consultant_rh.docx")
            self.generate_evaluations_report(all_submissions, output_path)
            self.stdout.write(f"Rapport Consultant RH généré : {output_path}")

        except Domain.DoesNotExist:
            self.stdout.write("Le domaine 'Consultant RH' n'existe pas")

    def generate_challenge_report(self, challenge, submissions, output_path):
        """Génère un rapport pour un challenge spécifique"""
        doc = Document()

        # Configuration de base du document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Ajouter un titre principal
        title = doc.add_heading(f'Rapport du challenge: {challenge.title}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Informations sur le challenge
        doc.add_paragraph(f"Domaine: {challenge.domain.name}")
        doc.add_paragraph(f"Description: {challenge.description}")
        doc.add_paragraph(f"Durée: {challenge.duration}")
        doc.add_paragraph(f"Nombre de questions: {challenge.questions.count()}")
        doc.add_paragraph(f"Nombre de soumissions: {submissions.count()}")

        # Statistiques des résultats
        doc.add_heading('Statistiques des résultats', level=1)

        results = [s.result * 100 if s.result else 0 for s in submissions]
        if results:
            avg_result = sum(results) / len(results)
            max_result = max(results) if results else 0
            min_result = min(results) if results else 0

            stats_table = doc.add_table(rows=1, cols=2)
            stats_table.style = 'Table Grid'
            row = stats_table.rows[0].cells
            row[0].text = 'Statistique'
            row[1].text = 'Valeur'

            row = stats_table.add_row().cells
            row[0].text = 'Moyenne des résultats'
            row[1].text = f"{avg_result:.2f}%"

            row = stats_table.add_row().cells
            row[0].text = 'Meilleur résultat'
            row[1].text = f"{max_result:.2f}%"

            row = stats_table.add_row().cells
            row[0].text = 'Résultat le plus bas'
            row[1].text = f"{min_result:.2f}%"

        # Détails des soumissions
        doc.add_heading('Détails des soumissions', level=1)

        for submission in submissions:
            self.add_submission_details(doc, submission)

        doc.save(output_path)
        return output_path

    def add_submission_details(self, doc, submission):
        """Ajoute les détails d'une soumission au document"""
        doc.add_heading(f'Candidat: {submission.candidate.get_full_name()}', level=2)

        # Table d'informations sur le candidat
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        row_cells = table.rows[0].cells
        row_cells[0].text = f'Email: {submission.candidate.email}'
        row_cells[1].text = f'Date: {submission.submitted_at.strftime("%d/%m/%Y %H:%M")}'

        row = table.add_row().cells
        row[0].text = f'Résultat: {submission.result_percent:.2f}%'
        row[1].text = f'Status: {submission.get_status_display()}'

        # Expérience du candidat
        if hasattr(submission.candidate, 'experience_level'):
            row = table.add_row().cells
            row[0].text = f'Niveau d\'expérience: {submission.candidate.get_experience_level_display()}'
            row[1].text = f'Années d\'expérience: {submission.candidate.experience}'

        doc.add_paragraph()  # Espacement

        # Ajouter les questions et réponses
        doc.add_heading('Questions et réponses', level=3)

        for i, answer in enumerate(submission.answers.all(), 1):
            # Titre de la question
            q_title = doc.add_paragraph()
            q_title.add_run(f'Question {i}: {answer.question.title}').bold = True

            # Informations sur la question
            q_info = doc.add_paragraph()
            q_info.add_run(f'Catégorie: {answer.question.category.name}\n')
            q_info.add_run(f'Type: {answer.question.get_question_type_display()}\n')
            q_info.add_run(f'Niveau: {answer.question.get_level_display()}\n')

            # Section réponse
            response_p = doc.add_paragraph()
            response_p.add_run('Réponse: ').bold = True

            if answer.question.is_open_answer:
                response_p.add_run(f'{answer.text}\n')

                # Afficher si la réponse est correcte
                result_p = doc.add_paragraph()
                result_run = result_p.add_run(f'Évaluation: {"Correct" if answer.is_correct else "Incorrect"}')
                result_run.bold = True
                result_run.font.color.rgb = RGBColor(0, 128, 0) if answer.is_correct else RGBColor(255, 0, 0)

            else:  # Questions à choix
                # Tableau des choix
                choices_table = doc.add_table(rows=1, cols=3)
                choices_table.style = 'Table Grid'
                choices_header = choices_table.rows[0].cells
                choices_header[0].text = "Choix"
                choices_header[1].text = "Sélectionné"
                choices_header[2].text = "Correct"

                # Remplir les choix
                for choice in answer.question.choices.all():
                    selected = choice in answer.selected_choices.all()
                    row = choices_table.add_row().cells
                    row[0].text = choice.text
                    row[1].text = "✓" if selected else ""
                    row[2].text = "✓" if choice.is_correct else ""

            doc.add_paragraph()  # Espacement entre questions

        doc.add_page_break()  # Séparer les candidats

    def generate_evaluations_report(self, submissions, output_path):
        """
        Génère un rapport Word pour une liste d'évaluations.

        Args:
            submissions: Liste des objets Submission
            output_path: Chemin de sortie pour le document Word
        """
        # Créer un document Word
        doc = Document()

        # Configuration de base du document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Ajouter un titre principal
        title = doc.add_heading('Rapport d\'évaluations', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Pour chaque évaluation
        for submission in submissions:
            # Ajouter les informations du candidat
            doc.add_heading(f'Candidat: {submission.candidate.get_full_name()}', level=1)

            # Table d'informations sur le candidat
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.autofit = True

            # Première ligne avec email et date
            row_cells = table.rows[0].cells
            row_cells[0].text = f'Email: {submission.candidate.email}'
            row_cells[1].text = f'Date: {submission.submitted_at.strftime("%d/%m/%Y %H:%M")}'

            # Deuxième ligne avec challenge et domaine
            row = table.add_row().cells
            row[0].text = f'Challenge: {submission.challenge.title}'
            row[1].text = f'Domaine: {submission.challenge.domain.name}'

            # Troisième ligne avec résultat
            row = table.add_row().cells
            row[0].text = f'Résultat: {submission.result_percent:.2f}%'
            row[1].text = f'Status: {submission.get_status_display()}'

            doc.add_paragraph()  # Espacement

            # Ajouter les questions et réponses
            doc.add_heading('Questions et réponses', level=2)

            for i, answer in enumerate(submission.answers.all(), 1):
                # Titre de la question
                q_title = doc.add_paragraph()
                q_title.add_run(f'Question {i}: {answer.question.title}').bold = True

                # Informations sur la question
                q_info = doc.add_paragraph()
                q_info.add_run(f'Catégorie: {answer.question.category.name}\n')
                q_info.add_run(f'Type: {answer.question.get_question_type_display()}\n')
                q_info.add_run(f'Niveau: {answer.question.get_level_display()}\n')

                # Section réponse
                response_p = doc.add_paragraph()
                response_p.add_run('Réponse: ').bold = True

                if answer.question.is_open_answer:
                    response_p.add_run(f'{answer.text}\n')

                    # Afficher si la réponse est correcte
                    result_p = doc.add_paragraph()
                    result_run = result_p.add_run(f'Évaluation: {"Correct" if answer.is_correct else "Incorrect"}')
                    result_run.bold = True
                    result_run.font.color.rgb = RGBColor(0, 128, 0) if answer.is_correct else RGBColor(255, 0, 0)

                else:  # Questions à choix
                    # Tableau des choix
                    choices_table = doc.add_table(rows=1, cols=3)
                    choices_table.style = 'Table Grid'
                    choices_header = choices_table.rows[0].cells
                    choices_header[0].text = "Choix"
                    choices_header[1].text = "Sélectionné"
                    choices_header[2].text = "Correct"

                    # Remplir les choix
                    for choice in answer.question.choices.all():
                        selected = choice in answer.selected_choices.all()
                        row = choices_table.add_row().cells
                        row[0].text = choice.text
                        row[1].text = "✓" if selected else ""
                        row[2].text = "✓" if choice.is_correct else ""

                doc.add_paragraph()  # Espacement entre questions

            # Séparateur entre candidats
            doc.add_page_break()

        # Enregistrer le document
        doc.save(output_path)
        return output_path