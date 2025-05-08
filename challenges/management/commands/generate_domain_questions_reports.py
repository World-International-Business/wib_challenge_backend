import os

from django.core.management.base import BaseCommand
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from questions.models import Domain, Question


class Command(BaseCommand):
    help = 'Génère un rapport Word des questions par domaine'

    def add_arguments(self, parser):
        parser.add_argument('--domain', type=str, required=True, help='Nom du domaine')
        parser.add_argument('--output-dir', type=str, default='reports', help='Répertoire de sortie')

    def handle(self, *args, **options):
        domain_name = options['domain']
        output_dir = options['output_dir']

        try:
            domain = Domain.objects.get(name=domain_name)
            self.stdout.write(f"Génération du rapport pour le domaine: {domain.name}")

            # Créer le répertoire de sortie s'il n'existe pas
            os.makedirs(output_dir, exist_ok=True)

            # Générer le rapport
            output_path = os.path.join(output_dir, f"questions_{domain.name.lower().replace(' ', '_')}.docx")
            self.generate_questions_report(domain, output_path)

            self.stdout.write(self.style.SUCCESS(f"Rapport généré avec succès: {output_path}"))

            # Teste spécifiquement pour Consultant RH
            if domain_name == "Consultant RH":
                self.stdout.write("Test effectué pour le domaine Consultant RH")

        except Domain.DoesNotExist:
            self.stderr.write(self.style.ERROR(f"Le domaine '{domain_name}' n'existe pas."))

    def generate_questions_report(self, domain, output_path):
        """
        Génère un rapport Word contenant toutes les questions d'un domaine.

        Args:
            domain: Objet Domain
            output_path: Chemin du fichier de sortie
        """
        # Créer le document Word
        doc = Document()

        # Configuration de base du document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Titre principal
        title = doc.add_heading(f'Questions du domaine: {domain.name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Informations générales
        total_questions = Question.objects.filter(category__domain=domain).count()
        open_answers = Question.objects.filter(
            category__domain=domain,
            question_type=Question.QuestionType.OPEN_ANSWER
        ).count()
        mcqs = Question.objects.filter(
            category__domain=domain,
            question_type=Question.QuestionType.MULTIPLE_CHOICE
        ).count()
        ucqs = Question.objects.filter(
            category__domain=domain,
            question_type=Question.QuestionType.UNIQUE_CHOICE
        ).count()

        doc.add_paragraph(f"Nombre total de questions: {total_questions}")
        doc.add_paragraph(f"Questions à réponse ouverte: {open_answers}")
        doc.add_paragraph(f"Questions à choix multiples: {mcqs}")
        doc.add_paragraph(f"Questions à choix unique: {ucqs}")

        # Organiser les questions par catégorie
        categories = domain.categories.all()

        for category in categories:
            doc.add_heading(f'Catégorie: {category.name}', level=1)

            questions = Question.objects.filter(category=category)

            if not questions:
                doc.add_paragraph("Aucune question pour cette catégorie.")
                continue

            # Ajouter les questions
            for i, question in enumerate(questions, 1):
                self.add_question_to_doc(doc, question, i)

        # Enregistrer le document
        doc.save(output_path)
        return output_path

    def add_question_to_doc(self, doc, question, question_number):
        """
        Ajoute les détails d'une question au document.

        Args:
            doc: Objet Document
            question: Objet Question
            question_number: Numéro de la question
        """
        # Titre de la question avec son numéro
        q_title = doc.add_heading(f'Question {question_number}: {question.title}', level=2)

        # Table d'informations sur la question
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Première ligne avec type et niveau
        row = table.rows[0].cells
        row[0].text = f'Type: {question.get_question_type_display()}'
        row[1].text = f'Niveau: {question.get_level_display()}'

        # Deuxième ligne avec catégorie
        row = table.add_row().cells
        row[0].text = f'Catégorie: {question.category.name}'

        # Description
        if question.description:
            doc.add_paragraph("Description:").bold = True
            doc.add_paragraph(question.description)

        # Tags
        if question.tags.exists():
            tags_p = doc.add_paragraph("Tags: ")
            tags_p.add_run(", ".join([tag.name for tag in question.tags.all()]))

        # Si question à choix, ajouter les choix
        if not question.is_open_answer:
            doc.add_heading('Choix possibles:', level=3)

            # Tableau des choix
            choices_table = doc.add_table(rows=1, cols=2)
            choices_table.style = 'Table Grid'
            header_cells = choices_table.rows[0].cells
            header_cells[0].text = "Option"
            header_cells[1].text = "Correct"

            for choice in question.choices.all():
                row = choices_table.add_row().cells
                row[0].text = choice.text
                row[1].text = "✓" if choice.is_correct else "✗"

        # Séparateur entre questions
        doc.add_paragraph()
